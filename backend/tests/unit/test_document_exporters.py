"""Tests for DocxExporter, XlsxExporter, and document tool registration."""

import io
from typing import Any

import pytest
from docx import Document
from openpyxl import load_workbook

from app.core.artifacts.exporters.docx_exporter import (
    DocxExporter,
    docx_exporter,
)
from app.core.artifacts.exporters.xlsx_exporter import (
    XlsxExporter,
    xlsx_exporter,
)
from app.models.artifact import ArtifactType

# ------------------------------------------------------------------ #
# DocxExporter
# ------------------------------------------------------------------ #


def test_docx_valuation_memo() -> None:
    """DocxExporter should produce valid .docx for valuation_memo."""
    content: dict[str, Any] = {
        "methodology": "DCF with 10% discount rate",
        "comparables": [
            {"company": "Acme", "revenue": "10M", "multiple": "5x"},
        ],
        "assumptions": [
            {"assumption": "Growth rate", "value": "20%"},
        ],
        "range_low": 5000000,
        "range_high": 15000000,
        "recommended": 10000000,
        "narrative": "Strong growth trajectory.",
    }
    result = docx_exporter.export(
        ArtifactType.VALUATION_MEMO, "Series A Valuation", content,
    )
    assert isinstance(result, bytes)
    assert len(result) > 0
    doc = Document(io.BytesIO(result))
    assert len(doc.paragraphs) > 0


def test_docx_research_brief() -> None:
    """DocxExporter should produce valid .docx for research_brief."""
    content: dict[str, Any] = {
        "summary": "Market analysis findings",
        "key_findings": ["Finding 1", "Finding 2"],
        "recommendations": ["Rec 1"],
        "sources": ["Source A"],
    }
    result = docx_exporter.export(
        ArtifactType.RESEARCH_BRIEF, "Market Research", content,
    )
    doc = Document(io.BytesIO(result))
    assert len(doc.paragraphs) > 0


def test_docx_board_memo() -> None:
    """DocxExporter should produce valid .docx for board_memo."""
    content: dict[str, Any] = {
        "subject": "Q4 Update",
        "executive_summary": "Strong quarter.",
        "key_updates": ["Revenue up 30%", "New hire VP Eng"],
        "decisions_needed": ["Approve Series B"],
    }
    result = docx_exporter.export(
        ArtifactType.BOARD_MEMO, "Board Memo Q4", content,
    )
    doc = Document(io.BytesIO(result))
    assert len(doc.paragraphs) > 0


def test_docx_pitch_narrative() -> None:
    """DocxExporter should produce valid .docx for pitch_narrative."""
    content: dict[str, Any] = {
        "hook": "Imagine a world...",
        "problem_story": "The problem is...",
        "solution_reveal": "We built...",
        "ask": "We're raising $5M.",
    }
    result = docx_exporter.export(
        ArtifactType.PITCH_NARRATIVE, "Our Story", content,
    )
    doc = Document(io.BytesIO(result))
    assert len(doc.paragraphs) > 0


def test_docx_dataroom() -> None:
    """DocxExporter should produce .docx for dataroom_structure."""
    content: dict[str, Any] = {
        "categories": [
            {
                "name": "Corporate",
                "required_docs": ["Articles", "Bylaws"],
                "uploaded_docs": ["Articles"],
                "completion_pct": 50.0,
            },
        ],
    }
    result = docx_exporter.export(
        ArtifactType.DATAROOM_STRUCTURE, "Dataroom", content,
    )
    doc = Document(io.BytesIO(result))
    assert len(doc.paragraphs) > 0


def test_docx_rejects_unsupported() -> None:
    """Should raise ValueError for unsupported types."""
    with pytest.raises(ValueError, match="does not support"):
        docx_exporter.export(
            ArtifactType.DECK_OUTLINE, "Test", {},
        )


# ------------------------------------------------------------------ #
# XlsxExporter
# ------------------------------------------------------------------ #


def test_xlsx_financial_model() -> None:
    """XlsxExporter should produce valid .xlsx for financial_model."""
    content: dict[str, Any] = {
        "revenue_projections": [
            {"year": 2024, "revenue": 500000},
            {"year": 2025, "revenue": 1200000},
        ],
        "cost_projections": [
            {"year": 2024, "costs": 400000},
        ],
        "runway_months": 18,
        "burn_rate": 50000.0,
        "unit_economics": {"CAC": 120, "LTV": 480},
        "funding_scenarios": [
            {"scenario": "Base", "amount": 2000000},
        ],
    }
    result = xlsx_exporter.export(
        ArtifactType.FINANCIAL_MODEL, "Financial Model", content,
    )
    assert isinstance(result, bytes)
    wb = load_workbook(io.BytesIO(result))
    assert "Revenue Projections" in wb.sheetnames
    assert "Summary" in wb.sheetnames


def test_xlsx_kpi_dashboard() -> None:
    """XlsxExporter should produce valid .xlsx for kpi_dashboard."""
    content: dict[str, Any] = {
        "metrics": [
            {
                "name": "MRR",
                "current_value": 50000.0,
                "target_value": 100000.0,
                "unit": "USD",
                "trend": "up",
                "category": "Revenue",
            },
        ],
    }
    result = xlsx_exporter.export(
        ArtifactType.KPI_DASHBOARD, "KPI Dashboard", content,
    )
    wb = load_workbook(io.BytesIO(result))
    assert "KPI Dashboard" in wb.sheetnames


def test_xlsx_rejects_unsupported() -> None:
    """Should raise ValueError for unsupported types."""
    with pytest.raises(ValueError, match="does not support"):
        xlsx_exporter.export(
            ArtifactType.LEAN_CANVAS, "Test", {},
        )


# ------------------------------------------------------------------ #
# Singletons
# ------------------------------------------------------------------ #


def test_exporter_singletons() -> None:
    assert isinstance(docx_exporter, DocxExporter)
    assert isinstance(xlsx_exporter, XlsxExporter)


# ------------------------------------------------------------------ #
# Tool registration
# ------------------------------------------------------------------ #


def test_document_tool_registration() -> None:
    """generate_document should be in the tool registry."""
    from app.core.tools.document_tools import register_document_tools
    from app.core.tools.registry import tool_registry

    register_document_tools()
    assert tool_registry.get_definition("generate_document") is not None


def test_agents_have_generate_document() -> None:
    """Key agents should have generate_document in their tool list."""
    from app.core.tools.registry import AGENT_TOOL_MAP

    expected_agents = [
        "valuation-strategist",
        "lean-modeler",
        "kpi-dashboard",
        "dataroom-concierge",
        "storyteller",
        "market-oracle",
        "venture-architect",
    ]
    for agent_id in expected_agents:
        assert "generate_document" in AGENT_TOOL_MAP[agent_id], (
            f"{agent_id} missing generate_document"
        )
