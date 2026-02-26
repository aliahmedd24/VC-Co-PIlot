"""DOCX exporter — converts text-heavy artifacts into Word documents.

Supports: valuation_memo, research_brief, board_memo, pitch_narrative,
dataroom_structure, lean_canvas.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor

from app.models.artifact import ArtifactType

# ---------------------------------------------------------------------------
# Design tokens
# ---------------------------------------------------------------------------

_ACCENT = RGBColor(0x00, 0x96, 0xFF)
_DARK = RGBColor(0x1A, 0x1A, 0x2E)
_MUTED = RGBColor(0x66, 0x66, 0x77)

_SUPPORTED_TYPES = {
    ArtifactType.VALUATION_MEMO,
    ArtifactType.RESEARCH_BRIEF,
    ArtifactType.BOARD_MEMO,
    ArtifactType.PITCH_NARRATIVE,
    ArtifactType.DATAROOM_STRUCTURE,
    ArtifactType.LEAN_CANVAS,
    ArtifactType.CUSTOM,
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _style_heading(paragraph: Any, level: int = 1) -> None:
    """Apply accent colour to heading runs."""
    for run in paragraph.runs:
        run.font.color.rgb = _DARK if level <= 1 else _ACCENT


def _add_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    font_size: int = 11,
    color: RGBColor = _DARK,
) -> Any:
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p


def _add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    """Add a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = str(value)


def _add_bullet_list(doc: Document, items: list[str]) -> None:
    """Add a bulleted list."""
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


# ---------------------------------------------------------------------------
# Per-type builders
# ---------------------------------------------------------------------------


def _build_valuation_memo(doc: Document, content: dict[str, Any]) -> None:
    """Build a valuation memo document."""
    if content.get("methodology"):
        doc.add_heading("Methodology", level=2)
        _add_paragraph(doc, content["methodology"])

    if content.get("comparables"):
        doc.add_heading("Comparable Analysis", level=2)
        headers = list(content["comparables"][0].keys())
        rows = [
            [str(c.get(h, "")) for h in headers]
            for c in content["comparables"]
        ]
        _add_table(doc, headers, rows)

    if content.get("assumptions"):
        doc.add_heading("Key Assumptions", level=2)
        headers = list(content["assumptions"][0].keys())
        rows = [
            [str(a.get(h, "")) for h in headers]
            for a in content["assumptions"]
        ]
        _add_table(doc, headers, rows)

    # Valuation range
    low = content.get("range_low")
    high = content.get("range_high")
    rec = content.get("recommended")
    if low is not None or high is not None:
        doc.add_heading("Valuation Range", level=2)
        parts = []
        if low is not None:
            parts.append(f"Low: ${low:,.0f}")
        if high is not None:
            parts.append(f"High: ${high:,.0f}")
        if rec is not None:
            parts.append(f"Recommended: ${rec:,.0f}")
        _add_paragraph(doc, " | ".join(parts), bold=True, font_size=14)

    if content.get("narrative"):
        doc.add_heading("Narrative", level=2)
        _add_paragraph(doc, content["narrative"])


def _build_research_brief(doc: Document, content: dict[str, Any]) -> None:
    """Build a research brief document."""
    if content.get("summary"):
        doc.add_heading("Executive Summary", level=2)
        _add_paragraph(doc, content["summary"])

    if content.get("methodology"):
        doc.add_heading("Methodology", level=2)
        _add_paragraph(doc, content["methodology"])

    if content.get("key_findings"):
        doc.add_heading("Key Findings", level=2)
        _add_bullet_list(doc, content["key_findings"])

    if content.get("recommendations"):
        doc.add_heading("Recommendations", level=2)
        _add_bullet_list(doc, content["recommendations"])

    if content.get("sources"):
        doc.add_heading("Sources", level=2)
        _add_bullet_list(doc, content["sources"])


def _build_board_memo(doc: Document, content: dict[str, Any]) -> None:
    """Build a board memo document."""
    if content.get("subject"):
        _add_paragraph(
            doc, f"RE: {content['subject']}",
            bold=True, font_size=13,
        )

    if content.get("executive_summary"):
        doc.add_heading("Executive Summary", level=2)
        _add_paragraph(doc, content["executive_summary"])

    if content.get("key_updates"):
        doc.add_heading("Key Updates", level=2)
        _add_bullet_list(doc, content["key_updates"])

    if content.get("financials_summary"):
        doc.add_heading("Financial Summary", level=2)
        _add_paragraph(doc, content["financials_summary"])

    if content.get("decisions_needed"):
        doc.add_heading("Decisions Required", level=2)
        _add_bullet_list(doc, content["decisions_needed"])

    if content.get("appendix"):
        doc.add_heading("Appendix", level=2)
        _add_bullet_list(doc, content["appendix"])


def _build_pitch_narrative(doc: Document, content: dict[str, Any]) -> None:
    """Build a pitch narrative document."""
    sections = [
        ("hook", "Opening Hook"),
        ("problem_story", "The Problem"),
        ("solution_reveal", "Our Solution"),
        ("traction_proof", "Traction & Proof"),
        ("market_opportunity", "Market Opportunity"),
        ("business_model", "Business Model"),
        ("team_story", "The Team"),
        ("ask", "The Ask"),
        ("vision", "Vision"),
    ]
    for key, heading in sections:
        text = content.get(key, "")
        if text:
            doc.add_heading(heading, level=2)
            _add_paragraph(doc, text)


def _build_dataroom(doc: Document, content: dict[str, Any]) -> None:
    """Build a dataroom checklist document."""
    categories = content.get("categories", [])
    if not categories:
        _add_paragraph(doc, "No categories defined.", italic=True)
        return

    for cat in categories:
        name = cat.get("name", "Unnamed")
        pct = cat.get("completion_pct", 0)
        doc.add_heading(f"{name} ({pct:.0f}% complete)", level=2)

        required = cat.get("required_docs", [])
        uploaded = set(cat.get("uploaded_docs", []))
        for d in required:
            status = "✅" if d in uploaded else "⬜"
            doc.add_paragraph(f"{status}  {d}", style="List Bullet")


def _build_lean_canvas(doc: Document, content: dict[str, Any]) -> None:
    """Build a lean canvas document."""
    sections = [
        ("problem", "Problem"),
        ("solution", "Solution"),
        ("unique_value_prop", "Unique Value Proposition"),
        ("unfair_advantage", "Unfair Advantage"),
        ("customer_segments", "Customer Segments"),
        ("key_metrics", "Key Metrics"),
        ("channels", "Channels"),
        ("cost_structure", "Cost Structure"),
        ("revenue_streams", "Revenue Streams"),
    ]
    for key, heading in sections:
        value = content.get(key)
        if not value:
            continue
        doc.add_heading(heading, level=2)
        if isinstance(value, list):
            _add_bullet_list(doc, value)
        else:
            _add_paragraph(doc, str(value))


def _build_custom(doc: Document, content: dict[str, Any]) -> None:
    """Build a generic document from custom content."""
    if content.get("body"):
        _add_paragraph(doc, content["body"])

    for section in content.get("sections", []):
        if section.get("title"):
            doc.add_heading(section["title"], level=2)
        if section.get("body"):
            _add_paragraph(doc, section["body"])


_BUILDERS: dict[ArtifactType, Any] = {
    ArtifactType.VALUATION_MEMO: _build_valuation_memo,
    ArtifactType.RESEARCH_BRIEF: _build_research_brief,
    ArtifactType.BOARD_MEMO: _build_board_memo,
    ArtifactType.PITCH_NARRATIVE: _build_pitch_narrative,
    ArtifactType.DATAROOM_STRUCTURE: _build_dataroom,
    ArtifactType.LEAN_CANVAS: _build_lean_canvas,
    ArtifactType.CUSTOM: _build_custom,
}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


class DocxExporter:
    """Export text-heavy artifacts to .docx bytes."""

    def export(
        self,
        artifact_type: ArtifactType,
        title: str,
        content: dict[str, Any],
    ) -> bytes:
        """Generate DOCX bytes from artifact content."""
        if artifact_type not in _SUPPORTED_TYPES:
            msg = (
                f"DocxExporter does not support "
                f"{artifact_type.value}"
            )
            raise ValueError(msg)

        doc = Document()

        # Title
        heading = doc.add_heading(title, level=0)
        _style_heading(heading, 0)

        # Subtitle with artifact type
        _add_paragraph(
            doc,
            artifact_type.value.replace("_", " ").title(),
            italic=True, font_size=12, color=_MUTED,
        )

        doc.add_paragraph()  # spacer

        # Delegate to type-specific builder
        builder = _BUILDERS.get(artifact_type, _build_custom)
        builder(doc, content)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


docx_exporter = DocxExporter()
